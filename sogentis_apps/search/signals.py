# search/signals.py
from __future__ import annotations

import logging
import mimetypes
import os
from threading import Lock
from typing import List, Tuple

from django.apps import apps
from django.conf import settings
from django.db.models.signals import post_delete, post_save

from .models import IndexedDocument

logger = logging.getLogger(__name__)

_registered: set[str] = set()
_lock = Lock()


# ------------------------------------------------------------
# Extraction texte (optionnelle)
# ------------------------------------------------------------
def extract_text_from_file(file_path: str) -> str:
    if not file_path or not os.path.exists(file_path):
        return ""

    mime_type, _ = mimetypes.guess_type(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    # PDF
    if mime_type == "application/pdf" or ext == ".pdf":
        try:
            import pdfplumber  # optional
            with pdfplumber.open(file_path) as pdf:
                return "\n".join([(p.extract_text() or "") for p in pdf.pages])
        except Exception as e:
            logger.debug("pdfplumber failed for %s: %s", file_path, e)

    # DOCX
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == ".docx":
        try:
            from docx import Document as DocxDocument  # optional
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            logger.debug("python-docx failed for %s: %s", file_path, e)

    # Text
    if (mime_type and mime_type.startswith("text/")) or ext in (".txt", ".csv", ".log"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
                return fh.read()
        except Exception as e:
            logger.debug("reading text file failed for %s: %s", file_path, e)

    # Images OCR (optionnel)
    if mime_type and mime_type.startswith("image/"):
        try:
            from PIL import Image  # optional
            import pytesseract     # optional
            img = Image.open(file_path)
            return pytesseract.image_to_string(img)
        except Exception as e:
            logger.debug("OCR failed for %s: %s", file_path, e)

    return ""


# ------------------------------------------------------------
# Modèles candidats
# ------------------------------------------------------------
def _parse_models_list(value) -> List[Tuple[str, str]]:
    out = []
    if not value:
        return out
    if isinstance(value, str):
        value = [value]
    for item in value:
        if isinstance(item, str) and "." in item:
            app_label, model_name = item.split(".", 1)
            out.append((app_label.strip(), model_name.strip()))
    return out


def get_candidate_models() -> List[Tuple[str, str]]:
    """
    1) SEARCH_SOURCE_MODELS = ["app.Model", ...]
    2) SEARCH_SOURCE_MODEL  = "app.Model"
    3) fallback list
    """
    models_list = _parse_models_list(getattr(settings, "SEARCH_SOURCE_MODELS", None))
    if models_list:
        return models_list

    single = getattr(settings, "SEARCH_SOURCE_MODEL", None)
    single_list = _parse_models_list(single)
    if single_list:
        return single_list

    return [
        ("social", "Document"),
        ("social", "Publication"),
        ("social", "Project"),
        ("social", "Engagement"),
        # ("social", "Donation"),  # active seulement si Donation a title/description
    ]


# ------------------------------------------------------------
# Mapping instance -> index
# ------------------------------------------------------------
def _safe_get(instance, name: str, default=""):
    try:
        return getattr(instance, name, default)
    except Exception:
        return default


def _file_info(instance):
    f = _safe_get(instance, "file", None)
    if not f:
        return "", ""

    try:
        file_url = getattr(f, "url", "") or str(f)
    except Exception:
        file_url = str(f)

    try:
        file_path = getattr(f, "path", "") or ""
    except Exception:
        file_path = ""

    return file_url, file_path


def _is_public(instance) -> bool:
    v = _safe_get(instance, "is_public", None)
    if v is None:
        return True
    return bool(v)


def _index_instance(sender, instance):
    title = (_safe_get(instance, "title", "") or "").strip()
    description = (_safe_get(instance, "description", "") or "").strip()

    body_text = (_safe_get(instance, "body", "") or "").strip()

    file_url, file_path = _file_info(instance)
    if not body_text and file_path:
        body_text = extract_text_from_file(file_path) or ""

    idx, _created = IndexedDocument.objects.get_or_create(
        source_app=sender._meta.app_label,
        source_model=sender._meta.model_name,
        object_id=_safe_get(instance, "id", None),
        defaults={
            "title": title,
            "description": description,
            "body": body_text,
            "file_url": file_url,
            "author": _safe_get(instance, "author", None),
            "is_public": _is_public(instance),
        },
    )

    changed = False

    updates = {
        "title": title,
        "description": description,
        "body": body_text,
        "file_url": file_url,
        "is_public": _is_public(instance),
    }

    for k, v in updates.items():
        v = v or ""
        if getattr(idx, k, "") != v:
            setattr(idx, k, v)
            changed = True

    author = _safe_get(instance, "author", None)
    if getattr(idx, "author", None) != author:
        idx.author = author
        changed = True

    if changed:
        idx.save()


def _remove_instance(sender, instance):
    IndexedDocument.objects.filter(
        source_app=sender._meta.app_label,
        source_model=sender._meta.model_name,
        object_id=_safe_get(instance, "id", None),
    ).delete()


def _on_save(sender, instance, **kwargs):
    try:
        _index_instance(sender, instance)
    except Exception as e:
        logger.exception(
            "Error indexing %s.%s id=%s: %s",
            sender._meta.app_label,
            sender.__name__,
            getattr(instance, "id", None),
            e,
        )


def _on_delete(sender, instance, **kwargs):
    try:
        _remove_instance(sender, instance)
    except Exception as e:
        logger.exception(
            "Error removing index %s.%s id=%s: %s",
            sender._meta.app_label,
            sender.__name__,
            getattr(instance, "id", None),
            e,
        )


def _connect_model(app_label: str, model_name: str) -> bool:
    key = f"{app_label}.{model_name}"

    with _lock:
        if key in _registered:
            return True

        try:
            Model = apps.get_model(app_label, model_name)
        except LookupError:
            logger.debug("apps.search: model %s.%s not found", app_label, model_name)
            return False

        post_save.connect(
            _on_save,
            sender=Model,
            dispatch_uid=f"search_index_save_{key}",
            weak=False,
        )
        post_delete.connect(
            _on_delete,
            sender=Model,
            dispatch_uid=f"search_index_delete_{key}",
            weak=False,
        )

        _registered.add(key)
        logger.info("apps.search: registered index handlers for %s.%s", app_label, model_name)
        return True


def ensure_handlers_registered():
    found = False
    for app_label, model_name in get_candidate_models():
        found = _connect_model(app_label, model_name) or found

    if not found:
        logger.info(
            "apps.search: no source model found. "
            "Set SEARCH_SOURCE_MODEL='social.Publication' "
            "or SEARCH_SOURCE_MODELS=['social.Publication', ...]"
        )


# Run at import (AppConfig.ready will import it as well)
try:
    ensure_handlers_registered()
except Exception:
    logger.debug("search.signals: deferred registration")






# # apps/search/signals.py
# """
# Signals for automatic indexing into IndexedDocument.

# - Dynamically identifies the model in app 'social' (Document, Publication, etc.)
#   or uses settings.SEARCH_SOURCE_MODEL = "app.Model".
# - Registers post_save / post_delete handlers for that model.
# - Extracts text content from files (PDF, DOCX, TXT, images).
# """

# import logging
# import os
# import mimetypes
# import threading

# from django.apps import apps
# from django.conf import settings
# from django.db.models.signals import post_save, post_delete
# from django.dispatch import receiver

# from .models import IndexedDocument

# logger = logging.getLogger(__name__)

# _registered_for = set()
# _lock = threading.Lock()


# def extract_text_from_file(file_path: str) -> str:
#     """Extract text from common document formats."""
#     if not file_path or not os.path.exists(file_path):
#         return ""

#     mime_type, _ = mimetypes.guess_type(file_path)
#     ext = os.path.splitext(file_path)[1].lower()

#     # PDF → pdfplumber
#     if mime_type == "application/pdf" or ext == ".pdf":
#         try:
#             import pdfplumber
#             with pdfplumber.open(file_path) as pdf:
#                 return "\n".join([p.extract_text() or "" for p in pdf.pages])
#         except Exception as e:
#             logger.debug("pdfplumber failed for %s: %s", file_path, e)

#     # DOCX → python-docx
#     if (
#         mime_type
#         == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         or ext == ".docx"
#     ):
#         try:
#             from docx import Document as DocxDocument
#             doc = DocxDocument(file_path)
#             return "\n".join([p.text for p in doc.paragraphs])
#         except Exception as e:
#             logger.debug("python-docx failed for %s: %s", file_path, e)

#     # Plain text
#     if (mime_type and mime_type.startswith("text/")) or ext in [".txt", ".csv", ".log"]:
#         try:
#             with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
#                 return fh.read()
#         except Exception as e:
#             logger.debug("reading text file failed for %s: %s", file_path, e)

#     # Images → OCR
#     if mime_type and mime_type.startswith("image/"):
#         try:
#             from PIL import Image
#             import pytesseract

#             img = Image.open(file_path)
#             return pytesseract.image_to_string(img)
#         except Exception as e:
#             logger.debug("OCR failed for %s: %s", file_path, e)

#     # Unsupported type
#     return ""


# def _get_candidate_models():
#     """Return list of candidate models to index."""
#     configured = getattr(settings, "SEARCH_SOURCE_MODEL", None)
#     if configured and isinstance(configured, str) and "." in configured:
#         app_label, model_name = configured.split(".", 1)
#         return [(app_label.strip(), model_name.strip())]

#     return [
#         ("social", "Document"),
#         ("social", "Publication"),
#         # ("social", "File"),
#         # ("social", "Attachment"),
#         # ("social", "Resource"),
#         ("social", "Project"),
#         ("social", "donation"),


#     ]


# def _ensure_handlers_for(app_label, model_name):
#     """Attach signals to a given model if available."""
#     key = f"{app_label}.{model_name}"
#     with _lock:
#         if key in _registered_for:
#             return

#         try:
#             SourceModel = apps.get_model(app_label, model_name)
#         except LookupError:
#             logger.debug("apps.search: model %s.%s not found", app_label, model_name)
#             return

#         @receiver(post_save, sender=SourceModel,  dispatch_uid=f"search_index_save_{key}")
#         def _index_source(sender, instance, **kwargs):
#             try:
#                 title = getattr(instance, "title", "") or ""
#                 description = getattr(instance, "description", "") or ""
#                 file_field = getattr(instance, "file", None)

#                 file_url = ""
#                 file_path = ""
#                 if file_field:
#                     file_url = getattr(file_field, "url", str(file_field))
#                     file_path = getattr(file_field, "path", "")

#                 body_text = getattr(instance, "body", "") or ""
#                 if not body_text and file_path:
#                     body_text = extract_text_from_file(file_path)

#                 idx, created_idx = IndexedDocument.objects.get_or_create(
#                     source_app=sender._meta.app_label,
#                     source_model=sender._meta.model_name,
#                     object_id=getattr(instance, "id", None),
#                     defaults={
#                         "title": title,
#                         "description": description,
#                         "body": body_text,
#                         "file_url": file_url,
#                         "author": getattr(instance, "author", None),
#                     },
#                 )

#                 changed = False
#                 for field_name, value in {
#                     "title": title,
#                     "description": description,
#                     "body": body_text,
#                     "file_url": file_url,
#                 }.items():
#                     if getattr(idx, field_name, "") != (value or ""):
#                         setattr(idx, field_name, value or "")
#                         changed = True

#                 if hasattr(instance, "author") and getattr(idx, "author", None) != instance.author:
#                     idx.author = instance.author
#                     changed = True

#                 if changed:
#                     idx.save()
#             except Exception as e:
#                 logger.exception(
#                     "Error indexing %s.%s id=%s: %s",
#                     app_label,
#                     model_name,
#                     getattr(instance, "id", None),
#                     e,
#                 )

#         @receiver(post_delete, sender=SourceModel, dispatch_uid=f"search_index_delete_{key}")
#         def _remove_index(sender, instance, **kwargs):
#             try:
#                 IndexedDocument.objects.filter(
#                     source_app=sender._meta.app_label,
#                     source_model=sender._meta.model_name,
#                     object_id=getattr(instance, "id", None),
#                 ).delete()
#             except Exception as e:
#                 logger.exception(
#                     "Error removing index for %s.%s id=%s: %s",
#                     app_label,
#                     model_name,
#                     getattr(instance, "id", None),
#                     e,
#                 )

#         _registered_for.add(key)
#         logger.info("apps.search: registered index handlers for %s.%s", app_label, model_name)


# def ensure_any_handlers_registered():
#     """Register signals for at least one candidate model."""
#     found = False
#     for app_label, model_name in _get_candidate_models():
#         _ensure_handlers_for(app_label, model_name)
#         if f"{app_label}.{model_name}" in _registered_for:
#             found = True
#     if not found:
#         logger.info(
#             "apps.search: no source model found. "
#             "Set SEARCH_SOURCE_MODEL in settings.py "
#             "(e.g. 'social.Publication')."
#         )


# # Try immediately (may be retried in AppConfig.ready)
# try:
#     ensure_any_handlers_registered()
# except Exception:
#     logger.debug("apps.search.signals: deferred registration until AppConfig.ready()")
